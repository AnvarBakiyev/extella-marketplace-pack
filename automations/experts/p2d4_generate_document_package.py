$extens("include.py")
include("import json", [])
include("import openpyxl", ["extella-pip install openpyxl"])
include("import docx", ["extella-pip install python-docx"])

def p2d4_generate_document_package(input_path: str = "", output_dir: str = "") -> dict:
    """Пакет документов из результата ИИ-анализа договоров (JSON от p2d4_evaluate_contract_batch):
    1) Реестр_рисков.xlsx — таблица: договор / критерий / статус / цитата / отклонение / статья ГК;
    2) Протокол_разногласий.docx — по каждому риску: текущая редакция, предлагаемая правка, обоснование;
    3) Сводка_руководителю.txt — короткий текст для отправки в Telegram/почту.
    Возвращает пути и ключевые цифры. Только чтение входа и запись в output_dir."""
    import json, os
    from pathlib import Path
    from datetime import datetime

    p = Path(str(input_path)).expanduser()
    if not str(input_path) or not p.exists():
        return {"status": "error", "message": "input_path не найден: " + str(input_path)}
    try:
        data = json.loads(p.read_text(encoding="utf-8"))
    except Exception as e:
        return {"status": "error", "message": "вход не JSON: " + str(e)[:120]}
    records = data.get("records") or []
    if not records:
        return {"status": "error", "message": "во входе нет записей с анализом"}

    out = Path(str(output_dir) or "/tmp/p2d4_docs").expanduser()
    out.mkdir(parents=True, exist_ok=True)
    stamp = datetime.now().strftime("%Y-%m-%d")

    def contract_name(i, rec):
        for k in ("file", "filename", "source", "source_file", "path", "name", "contract_name"):
            if rec.get(k):
                v = str(rec[k])
                return os.path.basename(v) if "/" in v or "\\" in v else v   # путь → имя файла
        # осмысленное имя из первых слов текста (пути отфильтровываем)
        for v in rec.values():
            s = str(v).strip()
            if len(s) > 40 and not s.startswith("/") and "\\" not in s[:10]:
                head = s.split("\n")[0][:60]
                if head:
                    return head
        return "Договор №" + str(i + 1)

    # ── 1. Реестр рисков (xlsx) ──
    import openpyxl
    from openpyxl.styles import Font, PatternFill, Alignment
    wb = openpyxl.Workbook(); ws = wb.active; ws.title = "Реестр рисков"
    headers = ["Договор", "Критерий", "Статус", "Цитата из договора", "Отклонение от стандарта", "Статья ГК", "Серьёзность"]
    ws.append(headers)
    for c in ws[1]:
        c.font = Font(bold=True, color="FFFFFF"); c.fill = PatternFill("solid", fgColor="8B5E2F")
        c.alignment = Alignment(vertical="center")
    fill_high = PatternFill("solid", fgColor="F8D7DA")
    fill_med = PatternFill("solid", fgColor="FFF3CD")
    n_rows = 0
    for i, rec in enumerate(records):
        a = rec.get("ai_analysis") or {}
        cname = contract_name(i, rec)
        dev_by_cond = {str(d.get("condition", "")): d for d in (a.get("deviations") or [])}
        for cr in (a.get("criteria") or []):
            if cr.get("status") == "ok":
                continue
            sev = ""
            dev_txt, law = "", ""
            for d in (a.get("deviations") or []):
                if d.get("condition", "").lower()[:6] in str(cr.get("criterion_name", cr.get("id", ""))).lower():
                    dev_txt = str(d.get("found", "")); law = str(d.get("law_ref", "")); sev = str(d.get("severity", ""))
                    break
            ws.append([cname, str(cr.get("criterion_name") or cr.get("id", "")), str(cr.get("status", "")),
                       str(cr.get("quote", ""))[:300], dev_txt[:200], law[:60], sev])
            n_rows += 1
            if sev == "high":
                for c in ws[ws.max_row]: c.fill = fill_high
            elif sev == "medium":
                for c in ws[ws.max_row]: c.fill = fill_med
        for d in (a.get("deviations") or []):
            ws.append([cname, str(d.get("condition", "")), "отклонение", "",
                       (str(d.get("found", "")) + " (норма: " + str(d.get("standard", "")) + ")")[:300],
                       str(d.get("law_ref", ""))[:60], str(d.get("severity", ""))])
            n_rows += 1
            if d.get("severity") == "high":
                for c in ws[ws.max_row]: c.fill = fill_high
    widths = [28, 26, 12, 48, 44, 14, 12]
    for j, w in enumerate(widths, 1):
        ws.column_dimensions[openpyxl.utils.get_column_letter(j)].width = w
    xlsx_path = str(out / ("Реестр_рисков_" + stamp + ".xlsx"))
    wb.save(xlsx_path)

    # ── 2. Протокол разногласий (docx) ──
    import docx
    doc = docx.Document()
    doc.add_heading("Протокол разногласий (черновик)", level=0)
    doc.add_paragraph("Подготовлен ИИ-агентом Extella " + stamp + ". Черновик — требует утверждения ответственным юристом; отправка контрагенту выполняется человеком.")
    total_edits = 0
    for i, rec in enumerate(records):
        a = rec.get("ai_analysis") or {}
        edits = a.get("suggested_edits") or []
        if not edits and not (a.get("deviations")):
            continue
        doc.add_heading(contract_name(i, rec), level=1)
        if a.get("summary"):
            doc.add_paragraph(str(a["summary"]))
        if a.get("risk_level"):
            doc.add_paragraph("Уровень риска: " + str(a["risk_level"]).upper())
        if edits:
            doc.add_heading("Предлагаемые правки", level=2)
            for e in edits:
                doc.add_paragraph(str(e), style="List Number")
                total_edits += 1
        devs = a.get("deviations") or []
        if devs:
            doc.add_heading("Отклонения от стандартных условий", level=2)
            for d in devs:
                line = (str(d.get("condition", "")) + ": в договоре — " + str(d.get("found", "не указано"))
                        + "; норма компании — " + str(d.get("standard", "")))
                if d.get("law_ref"):
                    line += "; основание — " + str(d["law_ref"])
                doc.add_paragraph(line, style="List Bullet")
    docx_path = str(out / ("Протокол_разногласий_" + stamp + ".docx"))
    doc.save(docx_path)

    # ── 3. Сводка руководителю (txt для Telegram/почты) ──
    n = len(records)
    highs = sum(1 for r in records if str((r.get("ai_analysis") or {}).get("risk_level", "")).lower() == "high")
    meds = sum(1 for r in records if str((r.get("ai_analysis") or {}).get("risk_level", "")).lower() == "medium")
    lines = ["📋 ИИ-анализ договоров — " + stamp,
             "Проверено документов: " + str(n),
             "⚠️ Высокий риск: " + str(highs) + " · Средний: " + str(meds),
             "Предложено правок: " + str(total_edits), ""]
    for i, rec in enumerate(records):
        a = rec.get("ai_analysis") or {}
        if str(a.get("risk_level", "")).lower() in ("high", "medium"):
            lines.append("• " + contract_name(i, rec)[:50] + " — риск " + str(a.get("risk_level", "")).upper()
                         + ": " + str(a.get("summary", ""))[:160])
    lines += ["", "Документы: Реестр рисков (Excel) + Протокол разногласий (Word) — готовы к согласованию.",
              "Ответьте «Согласовано» для финализации пакета."]
    summary_text = "\n".join(lines)
    txt_path = str(out / ("Сводка_руководителю_" + stamp + ".txt"))
    Path(txt_path).write_text(summary_text, encoding="utf-8")

    return {"status": "success", "output_dir": str(out), "registry_xlsx": xlsx_path,
            "protocol_docx": docx_path, "summary_txt": txt_path,
            "contracts": n, "high_risk": highs, "edits": total_edits, "risk_rows": n_rows}